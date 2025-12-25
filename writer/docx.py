from docx import Document

class WordWriter:
    def __init__(self, filename: str | None = None):
        """
        filename이 None이면 새 문서를 만들고,
        주어지면 해당 docx 파일을 엽니다.
        """
        if filename:
            self.doc = Document(filename)
            self.filename = filename
        else:
            self.doc = Document()
            self.filename = None

    def set_filename(self, filename: str):
        """저장할 파일 이름 설정."""
        self.filename = filename

    def add_title(self, text: str, level: int = 1):
        """제목(heading) 추가."""
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str):
        """일반 단락 추가."""
        return self.doc.add_paragraph(text)

    def add_mixed_paragraph(self, parts: list[dict]):
        """
        서로 다른 스타일을 가진 run들을 한 단락에 추가.
        parts 예시:
        [
            {"text": "굵게", "bold": True},
            {"text": " 보통글자", "bold": False}
        ]
        """
        p = self.doc.add_paragraph()
        for part in parts:
            run = p.add_run(part.get("text", ""))
            if part.get("bold"):
                run.bold = True
            if part.get("italic"):
                run.italic = True
        return p

    def save(self, filename: str | None = None):
        """문서 저장."""
        if filename is not None:
            self.filename = filename
        if not self.filename:
            raise ValueError("파일 이름이 설정되지 않았습니다.")
        self.doc
